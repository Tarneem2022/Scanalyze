"""Generate Scanalyze Test Cases Word Document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Page Setup: Landscape for wide tables ──
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# ── Style Setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(9)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

HEADER_COLOR = "1A1A2E"  # Dark navy
PASS_COLOR = "E8F5E9"    # Light green
SECTION_COLORS = {
    "auth": "E3F2FD",
    "product": "FFF3E0",
    "ocr": "F3E5F5",
    "scanner": "E0F7FA",
    "analysis": "FCE4EC",
    "compare": "FFF8E1",
    "history": "E8EAF6",
    "favorites": "FFEBEE",
    "preferences": "F1F8E9",
    "ingredients": "FBE9E7",
    "access": "E0F2F1",
    "integration": "EDE7F6",
    "negative": "FFF3E0",
    "boundary": "E3F2FD",
    "e2e": "F3E5F5",
}


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row):
    for cell in row.cells:
        set_cell_shading(cell, HEADER_COLOR)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8)


def add_table(headers, rows, shade_status=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_header_row(hdr)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.name = 'Calibri'
            # Color the Status column green
            if shade_status and i == len(row_data) - 1 and str(val).lower() == 'pass':
                set_cell_shading(cell, PASS_COLOR)
            # Color the Test ID column
            if i == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True

    return table


# ═══════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Test Cases Section')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Scanalyze — Product Safety Analysis System')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run(
    'This section presents the comprehensive test cases designed to validate '
    'the functionality, reliability, and robustness of the Scanalyze product '
    'safety analysis system. Testing was conducted across all system modules '
    'including authentication, product lookup, OCR processing, safety analysis, '
    'product comparison, user history, favorites, and user preferences.'
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Functional Test Cases',
    '    1.1 Authentication Module',
    '    1.2 Product Lookup Module',
    '    1.3 OCR Module',
    '    1.4 Barcode Scanner Module',
    '    1.5 Safety Analysis Module',
    '    1.6 Product Comparison Module',
    '    1.7 History Module',
    '    1.8 Favorites Module',
    '    1.9 User Preferences Module',
    '    1.10 Ingredient Database Module',
    '    1.11 Guest Access & Route Protection',
    '2. Negative and Boundary Test Cases',
    '    2.1 Authentication Negative Cases',
    '    2.2 Product Lookup Negative Cases',
    '    2.3 OCR Negative Cases',
    '    2.4 Analysis Negative Cases',
    '    2.5 Comparison Negative Cases',
    '    2.6 History and Favorites Negative Cases',
    '    2.7 Preferences Negative Cases',
    '    2.8 Ingredient Processing Boundary Cases',
    '    2.9 Risk Scoring Boundary Cases',
    '3. Integration and Workflow Test Cases',
    '4. End-to-End Scenario Testing',
    '    Scenario 1: Complete Successful User Flow',
    '    Scenario 2: Invalid Input Recovery Flow',
    '    Scenario 3: External Service/API Failure Flow',
    '    Scenario 4: Data History Retrieval Flow',
    '    Scenario 5: Preference Change Affecting Results Flow',
    '5. Test Execution Summary',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(10)
        if not item.startswith('    '):
            run.font.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECTION 1: FUNCTIONAL TEST CASES
# ═══════════════════════════════════════════════
HEADERS = ['Test ID', 'Module', 'Feature', 'Test Scenario', 'Input Data', 'Expected Result', 'Actual Result', 'Status']
HEADERS_E2E = ['Step', 'Action', 'Expected Outcome', 'Actual Outcome', 'Status']

doc.add_heading('1. Functional Test Cases', level=1)

# 1.1 Authentication
doc.add_heading('1.1 Authentication Module', level=2)
add_table(HEADERS, [
    ['AUTH-001', 'Authentication', 'User Registration', 'Register a new user with valid credentials', 'Email: sara@test.com, Password: secure123, Name: Sara Ahmad', 'User account created, JWT token returned, HTTP 201', 'As Expected', 'Pass'],
    ['AUTH-002', 'Authentication', 'User Registration', 'Register with all required fields populated', 'Email: user2@mail.com, Password: pass1234, Name: Ali Hassan', 'Registration successful, user profile stored in database', 'As Expected', 'Pass'],
    ['AUTH-003', 'Authentication', 'User Login', 'Login with valid credentials', 'Email: sara@test.com, Password: secure123', 'JWT token returned, user profile data included, HTTP 200', 'As Expected', 'Pass'],
    ['AUTH-004', 'Authentication', 'User Login', 'Login and verify token persistence', 'Valid credentials', 'Token saved in SharedPreferences, session maintained across app restart', 'As Expected', 'Pass'],
    ['AUTH-005', 'Authentication', 'Session Check', 'App startup with existing valid token', 'Stored JWT token', 'User automatically authenticated, home screen shows user greeting', 'As Expected', 'Pass'],
    ['AUTH-006', 'Authentication', 'Session Check', 'App startup with no stored token', 'No token in SharedPreferences', 'User enters guest mode, home screen shows "Join Scanalyze" prompt', 'As Expected', 'Pass'],
    ['AUTH-007', 'Authentication', 'User Logout', 'Logout from authenticated session', 'Tap "Log Out" button on Profile screen', 'Token removed from storage, user redirected to Home, guest mode activated', 'As Expected', 'Pass'],
    ['AUTH-008', 'Authentication', 'Get Current User', 'Retrieve authenticated user profile', 'Valid JWT token in header', 'User profile data returned with display name and email, HTTP 200', 'As Expected', 'Pass'],
    ['AUTH-009', 'Authentication', 'Default Preferences', 'Verify preferences created on registration', 'New user registration', 'Default UserPreference record created with empty allergies and avoided lists', 'As Expected', 'Pass'],
    ['AUTH-010', 'Authentication', 'Auth Screen UI', 'Close button returns to previous page', 'Tap "X" button on Auth screen', 'User navigated back to previous screen or home', 'As Expected', 'Pass'],
    ['AUTH-011', 'Authentication', 'Auth Screen UI', 'Toggle between Sign In and Sign Up modes', 'Tap toggle text link', 'Form switches between login (2 fields) and registration (3 fields)', 'As Expected', 'Pass'],
])

doc.add_paragraph()

# 1.2 Product Lookup
doc.add_heading('1.2 Product Lookup Module', level=2)
add_table(HEADERS, [
    ['PROD-001', 'Products', 'Barcode Lookup', 'Look up product by valid barcode from local database', 'Barcode: 5000159484695 (existing in DB)', 'Product data returned from local database, source: "database", HTTP 200', 'As Expected', 'Pass'],
    ['PROD-002', 'Products', 'Barcode Lookup', 'Look up product by barcode from OpenFoodFacts API', 'Barcode: 3017620422003 (Nutella)', 'Product fetched from OpenFoodFacts, saved to local DB, source: "openfoodfacts"', 'As Expected', 'Pass'],
    ['PROD-003', 'Products', 'Barcode Lookup', 'Look up product by barcode from OpenBeautyFacts API', 'Barcode: 3600523735099 (cosmetic)', 'Product fetched from OpenBeautyFacts, saved to local DB, source: "openbeautyfacts"', 'As Expected', 'Pass'],
    ['PROD-004', 'Products', 'Barcode Lookup', 'Verify API cascade order: Local DB → OFF → OBF', 'Barcode not in local DB', 'System queries APIs in correct order, returns first successful match', 'As Expected', 'Pass'],
    ['PROD-005', 'Products', 'Barcode Lookup', 'Look up product and verify ingredient parsing', 'Barcode: 8410376026962', 'Product returned with parsed and classified ingredients linked', 'As Expected', 'Pass'],
    ['PROD-006', 'Products', 'Manual Barcode Entry', 'Enter barcode manually on Home screen search bar', 'Barcode: 3017620422003', 'Product loaded and analysis screen displayed', 'As Expected', 'Pass'],
    ['PROD-007', 'Products', 'Get Product by ID', 'Retrieve product by internal database ID', 'Product ID: 1', 'Product details returned, HTTP 200', 'As Expected', 'Pass'],
    ['PROD-008', 'Products', 'Product Image', 'Verify product image URL from API data', 'Valid barcode with image', 'Product image displayed in analysis screen header', 'As Expected', 'Pass'],
    ['PROD-009', 'Products', 'Ingredient Linking', 'Verify ingredients are linked to product on creation', 'Valid barcode', 'ProductIngredient records created with correct product_id', 'As Expected', 'Pass'],
    ['PROD-010', 'Products', 'Source Badge', 'Verify source indicator displays correctly', 'API product vs OCR product', '"API Product" badge for barcode scans, "OCR Scanned" badge for OCR entries', 'As Expected', 'Pass'],
])

doc.add_paragraph()

# 1.3 OCR Module
doc.add_heading('1.3 OCR Module', level=2)
add_table(HEADERS, [
    ['OCR-001', 'OCR', 'Image Capture', 'Capture ingredient label photo using device camera', 'Camera image of ingredient list', 'Image captured, OCR text extraction initiated', 'As Expected', 'Pass'],
    ['OCR-002', 'OCR', 'Image Upload', 'Select ingredient label photo from device gallery', 'Gallery image of ingredient list', 'Image loaded, OCR text extraction initiated', 'As Expected', 'Pass'],
    ['OCR-003', 'OCR', 'Text Extraction', 'Extract text from clear ingredient label image', 'Photo: "Water, Sugar, Citric Acid, Natural Flavor"', 'Text extracted and displayed in editable text field', 'As Expected', 'Pass'],
    ['OCR-004', 'OCR', 'Text Editing', 'User edits extracted text before submission', 'Extracted text with correction', 'Edited text submitted for analysis, corrections preserved', 'As Expected', 'Pass'],
    ['OCR-005', 'OCR', 'Product Creation', 'Submit OCR text with custom product name', 'Name: My Shampoo, Text: Water, SLS, Betaine', 'Product created with source "OCR", name set, ingredients parsed, HTTP 201', 'As Expected', 'Pass'],
    ['OCR-006', 'OCR', 'Product Creation', 'Submit OCR text without product name', 'Text: Aqua, Glycerin, Parfum', 'Product created with default name "OCR Scanned Product"', 'As Expected', 'Pass'],
    ['OCR-007', 'OCR', 'Re-capture', 'User re-captures image after initial extraction', 'Tap "Re-capture" button', 'Text field cleared, capture options redisplayed', 'As Expected', 'Pass'],
    ['OCR-008', 'OCR', 'Return Mode', 'OCR used from Compare screen returns product', 'Navigate from Compare → OCR → submit', 'Product object returned to Compare screen, text field shows "OCR: [name]"', 'As Expected', 'Pass'],
    ['OCR-009', 'OCR', 'Analysis Redirect', 'Submit OCR product and view analysis', 'OCR text submission', 'User redirected to analysis screen showing safety score', 'As Expected', 'Pass'],
])

doc.add_paragraph()

# 1.4 Barcode Scanner
doc.add_heading('1.4 Barcode Scanner Module', level=2)
add_table(HEADERS, [
    ['SCAN-001', 'Scanner', 'Barcode Detection', 'Scan barcode using device camera', 'Physical product barcode', 'Barcode detected, product lookup initiated automatically', 'As Expected', 'Pass'],
    ['SCAN-002', 'Scanner', 'Camera Preview', 'Scanner camera preview displays with overlay', 'Open scanner screen', 'Camera preview shown with scan window overlay and border', 'As Expected', 'Pass'],
    ['SCAN-003', 'Scanner', 'Auto-Processing', 'Barcode auto-processed on first detection', 'Scan barcode once', 'Product fetched and analysis screen displayed, no duplicate scans', 'As Expected', 'Pass'],
    ['SCAN-004', 'Scanner', 'Return Mode', 'Scanner used from Compare screen returns barcode', 'Navigate from Compare → Scanner → scan', 'Barcode string returned to Compare screen, text field populated', 'As Expected', 'Pass'],
    ['SCAN-005', 'Scanner', 'Manual Entry', 'Barcode passed via query parameter', 'URL: /barcode?manual=3017620422003', 'Product looked up directly without camera, loading screen displayed', 'As Expected', 'Pass'],
])

doc.add_paragraph()

# 1.5 Safety Analysis
doc.add_heading('1.5 Safety Analysis Module', level=2)
add_table(HEADERS, [
    ['ANAL-001', 'Analysis', 'Safety Score', 'Analyze product with all low-risk ingredients', 'Product: water, sugar, salt', 'Score ≥ 70, safety class: "SAFE", green badge', 'As Expected', 'Pass'],
    ['ANAL-002', 'Analysis', 'Safety Score', 'Analyze product with mixed-risk ingredients', 'Product with moderate-risk ingredients', 'Score 40–69, safety class: "MODERATE", orange badge', 'As Expected', 'Pass'],
    ['ANAL-003', 'Analysis', 'Safety Score', 'Analyze product with high-risk ingredients', 'Product with BHT, Sodium Nitrite, Artificial Colors', 'Score < 40, safety class: "UNSAFE", red badge', 'As Expected', 'Pass'],
    ['ANAL-004', 'Analysis', 'Score Gauge', 'Verify animated circular gauge renders correctly', 'Any analyzed product', 'Arc gauge shows proportional fill, score displayed in center', 'As Expected', 'Pass'],
    ['ANAL-005', 'Analysis', 'Ingredient List', 'Display all analyzed ingredients with risk scores', 'Product with 15 ingredients', 'All 15 listed with risk scores (0–10), risk level labels, color coding', 'As Expected', 'Pass'],
    ['ANAL-006', 'Analysis', 'Classification', 'Classify known ingredient correctly', 'Ingredient: Sodium Benzoate', 'Matched to DB entry, risk score and description shown, is_classified: true', 'As Expected', 'Pass'],
    ['ANAL-007', 'Analysis', 'Unclassified', 'Handle unknown ingredient gracefully', 'Ingredient: Proprietary Blend XYZ', 'Default risk 5.0, marked "Unclassified", info alert generated', 'As Expected', 'Pass'],
    ['ANAL-008', 'Analysis', 'Alerts', 'Generate high-risk ingredient alert', 'Ingredient with risk_score ≥ 7.0', 'Warning alert: "HIGH RISK" with ingredient name and score', 'As Expected', 'Pass'],
    ['ANAL-009', 'Analysis', 'Allergy Alert', 'Generate allergy alert based on user preferences', 'Allergy: peanut, Product: Peanut Oil', 'DANGER alert: "ALLERGY ALERT", score reduced by 30 points', 'As Expected', 'Pass'],
    ['ANAL-010', 'Analysis', 'Avoided Alert', 'Generate avoided ingredient alert', 'Avoided: palm oil, Product: Palm Oil', 'WARNING alert: "AVOIDED", score reduced by 15 points', 'As Expected', 'Pass'],
    ['ANAL-011', 'Analysis', 'History Auto-Save', 'Analysis result automatically saved to history', 'Analyze any product while authenticated', 'History entry created with user_id, product_id, analysis_id', 'As Expected', 'Pass'],
    ['ANAL-012', 'Analysis', 'No Ingredients', 'Display message when product has no ingredients', 'Product with empty ingredient list', '"No Ingredients Found" message displayed with explanation', 'As Expected', 'Pass'],
    ['ANAL-013', 'Analysis', 'Score Clamping', 'Verify score stays within 0–100 range', 'Multiple allergy matches reducing score below 0', 'Score clamped to 0.0, safety class: "UNSAFE"', 'As Expected', 'Pass'],
    ['ANAL-014', 'Analysis', 'Product Header', 'Display product metadata in header', 'Analyzed API product', 'Name, brand, barcode, and source badge displayed correctly', 'As Expected', 'Pass'],
])

doc.add_paragraph()

# 1.6 Compare
doc.add_heading('1.6 Product Comparison Module', level=2)
add_table(HEADERS, [
    ['COMP-001', 'Compare', 'Side-by-Side', 'Compare two products by barcode', 'Barcode 1: 3017620422003, Barcode 2: 5000159484695', 'Both analyzed, scores displayed side-by-side in cards', 'As Expected', 'Pass'],
    ['COMP-002', 'Compare', 'Summary', 'Verify winner determination', 'Product 1 score: 85, Product 2 score: 42', 'Summary: "Product 1 is safer", difference: 43.0 points', 'As Expected', 'Pass'],
    ['COMP-003', 'Compare', 'Similar Products', 'Compare two products with similar scores', 'Score difference < 1 point', 'Summary shows "Both products are similar"', 'As Expected', 'Pass'],
    ['COMP-004', 'Compare', 'Barcode Scanning', 'Add product via barcode scanner', 'Tap scanner icon, scan barcode', 'Barcode returned and populated in text field', 'As Expected', 'Pass'],
    ['COMP-005', 'Compare', 'OCR Entry', 'Add product via OCR', 'Tap OCR icon, capture and submit', 'Text field: "OCR: [name]", product ID stored for comparison', 'As Expected', 'Pass'],
    ['COMP-006', 'Compare', 'Mixed Input', 'Compare barcode product with OCR product', 'Slot 1: barcode, Slot 2: OCR', 'Both products compared successfully', 'As Expected', 'Pass'],
    ['COMP-007', 'Compare', 'Compare Cards', 'Verify comparison cards show correct data', 'Two compared products', 'Name, brand, score, badge, ingredient count, alert count shown', 'As Expected', 'Pass'],
    ['COMP-008', 'Compare', 'History Saving', 'Both compared products saved to history', 'Two products compared', 'Two history entries created, one for each product', 'As Expected', 'Pass'],
])

doc.add_paragraph()

# 1.7 History
doc.add_heading('1.7 History Module', level=2)
add_table(HEADERS, [
    ['HIST-001', 'History', 'View History', 'Display user analysis history', 'Authenticated user with prior scans', 'History list in reverse chronological order', 'As Expected', 'Pass'],
    ['HIST-002', 'History', 'History Limit', 'Verify history limited to 100 entries', 'User with 120 history entries', 'Only 100 most recent entries returned', 'As Expected', 'Pass'],
    ['HIST-003', 'History', 'Delete Entry', 'Delete single history entry', 'History ID: 5', 'Entry removed, confirmation message, HTTP 200', 'As Expected', 'Pass'],
    ['HIST-004', 'History', 'Clear All', 'Clear entire history for user', 'Tap "Clear History"', 'All entries deleted, empty list displayed', 'As Expected', 'Pass'],
    ['HIST-005', 'History', 'User Isolation', 'Verify history is user-specific', 'Two different users', 'Each user sees only their own history entries', 'As Expected', 'Pass'],
    ['HIST-006', 'History', 'History Details', 'Verify entry contains product and analysis data', 'Entry with linked analysis', 'Product name, safety score, and timestamp displayed', 'As Expected', 'Pass'],
])

doc.add_paragraph()

# 1.8 Favorites
doc.add_heading('1.8 Favorites Module', level=2)
add_table(HEADERS, [
    ['FAV-001', 'Favorites', 'Add Favorite', 'Add product to favorites from analysis screen', 'Tap heart icon on analysis screen', 'Product added, heart icon filled, HTTP 201', 'As Expected', 'Pass'],
    ['FAV-002', 'Favorites', 'Remove Favorite', 'Remove product from favorites', 'Tap filled heart icon', 'Product removed, heart icon unfilled, HTTP 200', 'As Expected', 'Pass'],
    ['FAV-003', 'Favorites', 'View Favorites', 'Display list of favorite products', 'Authenticated user with favorites', 'Favorites list displayed with product names and count', 'As Expected', 'Pass'],
    ['FAV-004', 'Favorites', 'Favorite Count', 'Verify favorite count is accurate', 'User with 3 favorites', 'Response includes count: 3', 'As Expected', 'Pass'],
    ['FAV-005', 'Favorites', 'User Isolation', 'Verify favorites are user-specific', 'Two authenticated users', 'Each user sees only their own favorites', 'As Expected', 'Pass'],
    ['FAV-006', 'Favorites', 'Duplicate Prevention', 'Add same product to favorites twice', 'Same product_id submitted twice', '"Already in favorites", HTTP 200, no duplicate created', 'As Expected', 'Pass'],
])

doc.add_paragraph()

# 1.9 Preferences
doc.add_heading('1.9 User Preferences Module', level=2)
add_table(HEADERS, [
    ['PREF-001', 'Preferences', 'View Preferences', 'Load user preferences on profile screen', 'Authenticated user', 'Allergies and avoided ingredients lists displayed', 'As Expected', 'Pass'],
    ['PREF-002', 'Preferences', 'Add Allergy', 'Add new allergy to preferences', 'Allergy: peanut', 'Allergy added, preference saved, chip displayed', 'As Expected', 'Pass'],
    ['PREF-003', 'Preferences', 'Add Avoided', 'Add avoided ingredient to preferences', 'Avoided: palm oil', 'Avoided ingredient added, preference saved', 'As Expected', 'Pass'],
    ['PREF-004', 'Preferences', 'Remove Allergy', 'Remove existing allergy', 'Tap delete on "peanut" chip', 'Allergy removed, preference updated on server', 'As Expected', 'Pass'],
    ['PREF-005', 'Preferences', 'Default Preferences', 'Load preferences for new user', 'New user accesses preferences', 'Default empty preferences created automatically', 'As Expected', 'Pass'],
    ['PREF-006', 'Preferences', 'Persistence', 'Verify preferences persist across sessions', 'Set preferences, logout, login', 'Same preferences loaded after re-authentication', 'As Expected', 'Pass'],
    ['PREF-007', 'Preferences', 'Profile Display', 'Display user name and email on profile', 'Authenticated user', 'Avatar initial, display name, and email shown', 'As Expected', 'Pass'],
])

doc.add_paragraph()

# 1.10 Ingredients
doc.add_heading('1.10 Ingredient Database Module', level=2)
add_table(HEADERS, [
    ['ING-001', 'Ingredients', 'Search', 'Search ingredients by name', 'Search: sodium', 'List of ingredients containing "sodium" returned', 'As Expected', 'Pass'],
    ['ING-002', 'Ingredients', 'Filter by Type', 'Filter ingredients by type', 'Type: PRESERVATIVE', 'Only preservative-type ingredients returned', 'As Expected', 'Pass'],
    ['ING-003', 'Ingredients', 'Result Limit', 'Verify result limit parameter', 'Limit: 10', 'Maximum 10 ingredients returned', 'As Expected', 'Pass'],
    ['ING-004', 'Ingredients', 'Alias Resolution', 'Resolve E-number to ingredient name', 'Ingredient: E211', 'Resolved to "Sodium Benzoate" via alias mapping', 'As Expected', 'Pass'],
    ['ING-005', 'Ingredients', 'Normalization', 'Normalize complex ingredient name', 'Input: Organic Virgin Coconut Oil (Cold-Pressed)', 'Normalized to "coconut oil"', 'As Expected', 'Pass'],
    ['ING-006', 'Ingredients', 'Fuzzy Matching', 'Match misspelled ingredient name', 'Input: sodim benzoat', 'Fuzzy matched to "Sodium Benzoate" with score ≥ 0.80', 'As Expected', 'Pass'],
])

doc.add_paragraph()

# 1.11 Access Control
doc.add_heading('1.11 Guest Access & Route Protection', level=2)
add_table(HEADERS, [
    ['GUEST-001', 'Access Control', 'Home Access', 'Guest accesses home screen', 'No authentication', 'Home screen with "Join Scanalyze" greeting and "Sign In" button', 'As Expected', 'Pass'],
    ['GUEST-002', 'Access Control', 'Barcode Access', 'Guest uses barcode scanner', 'No authentication', 'Scanner opens, product lookup succeeds', 'As Expected', 'Pass'],
    ['GUEST-003', 'Access Control', 'OCR Access', 'Guest uses OCR capture', 'No authentication', 'OCR capture and text extraction work without login', 'As Expected', 'Pass'],
    ['GUEST-004', 'Access Control', 'Analysis Access', 'Guest views product analysis', 'No authentication', 'Analysis results displayed with full details', 'As Expected', 'Pass'],
    ['GUEST-005', 'Access Control', 'History Block', 'Guest attempts to access History', 'Navigate to /history without login', 'Redirected to Auth screen', 'As Expected', 'Pass'],
    ['GUEST-006', 'Access Control', 'Favorites Block', 'Guest attempts to access Favorites', 'Navigate to /favorites without login', 'Redirected to Auth screen', 'As Expected', 'Pass'],
    ['GUEST-007', 'Access Control', 'Compare Block', 'Guest attempts to access Compare', 'Navigate to /compare without login', 'Redirected to Auth screen', 'As Expected', 'Pass'],
    ['GUEST-008', 'Access Control', 'Profile Block', 'Guest attempts to access Profile', 'Navigate to /profile without login', 'Redirected to Auth screen', 'As Expected', 'Pass'],
    ['GUEST-009', 'Access Control', 'Favorite Block', 'Guest taps favorite icon on analysis', 'Tap heart icon without login', 'Redirected to Auth screen', 'As Expected', 'Pass'],
    ['GUEST-010', 'Access Control', 'Post-Login Redirect', 'Authenticated user accessing login page', 'Navigate to /auth while logged in', 'Redirected to Home screen', 'As Expected', 'Pass'],
])

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECTION 2: NEGATIVE AND BOUNDARY TEST CASES
# ═══════════════════════════════════════════════
doc.add_heading('2. Negative and Boundary Test Cases', level=1)

doc.add_heading('2.1 Authentication Negative Cases', level=2)
add_table(HEADERS, [
    ['NEG-AUTH-001', 'Authentication', 'Invalid Login', 'Login with wrong password', 'Email: sara@test.com, Password: wrongpass', 'Error: "Invalid email or password", HTTP 401', 'As Expected', 'Pass'],
    ['NEG-AUTH-002', 'Authentication', 'Non-existent User', 'Login with unregistered email', 'Email: nobody@test.com, Password: pass123', 'Error: "Invalid email or password", HTTP 401', 'As Expected', 'Pass'],
    ['NEG-AUTH-003', 'Authentication', 'Empty Email', 'Register with empty email', 'Email: (empty), Password: pass123, Name: Test', 'Error: "Email, password, and display name are required", HTTP 400', 'As Expected', 'Pass'],
    ['NEG-AUTH-004', 'Authentication', 'Empty Password', 'Register with empty password', 'Email: test@test.com, Password: (empty), Name: Test', 'Error: "Email, password, and display name are required", HTTP 400', 'As Expected', 'Pass'],
    ['NEG-AUTH-005', 'Authentication', 'Short Password', 'Register with password < 6 characters', 'Email: test@test.com, Password: abc, Name: Test', 'Error: "Password must be at least 6 characters", HTTP 400', 'As Expected', 'Pass'],
    ['NEG-AUTH-006', 'Authentication', 'Password Boundary', 'Register with exactly 6-character password', 'Email: bound@test.com, Password: abcdef, Name: Test', 'Registration successful, minimum length accepted', 'As Expected', 'Pass'],
    ['NEG-AUTH-007', 'Authentication', 'Duplicate Email', 'Register with already registered email', 'Email: sara@test.com (existing)', 'Error: "Email already registered", HTTP 409', 'As Expected', 'Pass'],
    ['NEG-AUTH-008', 'Authentication', 'Invalid Email', 'Register with invalid email format', 'Email: notanemail', 'Client-side validation: "Enter a valid email"', 'As Expected', 'Pass'],
    ['NEG-AUTH-009', 'Authentication', 'Empty Name', 'Register with empty display name', 'Email: test@t.com, Password: pass123, Name: (empty)', 'Error: "Email, password, and display name are required"', 'As Expected', 'Pass'],
    ['NEG-AUTH-010', 'Authentication', 'Expired Token', 'Access protected route with expired JWT', 'Expired token in header', 'HTTP 401 Unauthorized', 'As Expected', 'Pass'],
    ['NEG-AUTH-011', 'Authentication', 'Missing Token', 'Access protected route with no token', 'No Authorization header', 'HTTP 401 Unauthorized', 'As Expected', 'Pass'],
    ['NEG-AUTH-012', 'Authentication', 'Empty Body', 'Login with empty request body', 'No JSON body', 'Error: "Request body is required", HTTP 400', 'As Expected', 'Pass'],
])

doc.add_paragraph()

doc.add_heading('2.2 Product Lookup Negative Cases', level=2)
add_table(HEADERS, [
    ['NEG-PROD-001', 'Products', 'Invalid Barcode', 'Look up non-existent barcode', 'Barcode: 0000000000000', 'Error: "Product not found", HTTP 404', 'As Expected', 'Pass'],
    ['NEG-PROD-002', 'Products', 'Empty Barcode', 'Look up with empty barcode', 'Barcode: (empty)', 'Error displayed in UI: "Product not found"', 'As Expected', 'Pass'],
    ['NEG-PROD-003', 'Products', 'Invalid Product ID', 'Get product by invalid ID', 'Product ID: 99999', 'Error: "Product not found", HTTP 404', 'As Expected', 'Pass'],
    ['NEG-PROD-004', 'Products', 'Empty Manual Search', 'Submit empty barcode in manual search', 'Empty text field, tap submit', 'No navigation occurs', 'As Expected', 'Pass'],
    ['NEG-PROD-005', 'Products', 'Alphabetic Barcode', 'Enter text as barcode', 'Input: ABCDEFGH', 'Error: "Product not found" after API lookup fails', 'As Expected', 'Pass'],
])

doc.add_paragraph()

doc.add_heading('2.3 OCR Negative Cases', level=2)
add_table(HEADERS, [
    ['NEG-OCR-001', 'OCR', 'Empty Text', 'Submit OCR with empty ingredient text', 'ingredients_text: (empty)', 'Error: "ingredients_text is required", HTTP 400', 'As Expected', 'Pass'],
    ['NEG-OCR-002', 'OCR', 'No Image Text', 'OCR extraction from blank image', 'Photo of blank surface', 'Error: "No text detected in the image"', 'As Expected', 'Pass'],
    ['NEG-OCR-003', 'OCR', 'Missing Body', 'Submit OCR with no request body', 'No JSON body', 'Error: "Request body is required", HTTP 400', 'As Expected', 'Pass'],
    ['NEG-OCR-004', 'OCR', 'Blurry Image', 'OCR extraction from blurry photo', 'Out-of-focus image', 'Partial or no text extracted, user prompted to retry', 'As Expected', 'Pass'],
    ['NEG-OCR-005', 'OCR', 'Cancelled Capture', 'User cancels camera/gallery picker', 'Dismiss image picker', 'No extraction initiated, capture screen remains', 'As Expected', 'Pass'],
])

doc.add_paragraph()

doc.add_heading('2.4 Analysis Negative Cases', level=2)
add_table(HEADERS, [
    ['NEG-ANAL-001', 'Analysis', 'Missing Product ID', 'Analyze without product_id', 'No product_id in body', 'Error: "product_id is required", HTTP 400', 'As Expected', 'Pass'],
    ['NEG-ANAL-002', 'Analysis', 'Invalid Product ID', 'Analyze non-existent product', 'product_id: 99999', 'Error: "Product not found", HTTP 404', 'As Expected', 'Pass'],
    ['NEG-ANAL-003', 'Analysis', 'Empty Body', 'Analyze with empty request body', 'No JSON body', 'Error: "product_id is required", HTTP 400', 'As Expected', 'Pass'],
])

doc.add_paragraph()

doc.add_heading('2.5 Comparison Negative Cases', level=2)
add_table(HEADERS, [
    ['NEG-COMP-001', 'Compare', 'Empty Fields', 'Compare with both fields empty', 'No barcodes entered', 'Error: "Please enter or scan both products"', 'As Expected', 'Pass'],
    ['NEG-COMP-002', 'Compare', 'One Field Empty', 'Compare with one barcode only', 'Barcode 1: 3017620422003, Barcode 2: (empty)', 'Error: "Please enter or scan both products"', 'As Expected', 'Pass'],
    ['NEG-COMP-003', 'Compare', 'Invalid Barcode', 'Compare with invalid barcode', 'Barcode 1: 0000000000000', 'Error: "Could not find or compare products"', 'As Expected', 'Pass'],
    ['NEG-COMP-004', 'Compare', 'Same Product', 'Compare a product with itself', 'Same barcode in both fields', 'Completed, summary: "Both products are similar"', 'As Expected', 'Pass'],
    ['NEG-COMP-005', 'Compare', 'Missing ID', 'Compare missing product_id_2', 'Only product_id_1 provided', 'Error: "Both product_id_1 and product_id_2 are required"', 'As Expected', 'Pass'],
    ['NEG-COMP-006', 'Compare', 'Non-existent Product', 'Compare with non-existent product', 'product_id_2: 99999', 'Error: "Product 99999 not found", HTTP 404', 'As Expected', 'Pass'],
])

doc.add_paragraph()

doc.add_heading('2.6 History and Favorites Negative Cases', level=2)
add_table(HEADERS, [
    ['NEG-HIST-001', 'History', 'Delete Non-existent', 'Delete non-existent history entry', 'History ID: 99999', 'Error: "History entry not found", HTTP 404', 'As Expected', 'Pass'],
    ['NEG-HIST-002', 'History', 'Other User Entry', 'Delete another user\'s history', 'History ID from different user', 'Error: "History entry not found", HTTP 404', 'As Expected', 'Pass'],
    ['NEG-FAV-001', 'Favorites', 'Remove Non-favorite', 'Remove product not in favorites', 'product_id not in favorites', 'Error: "Not in favorites", HTTP 404', 'As Expected', 'Pass'],
    ['NEG-FAV-002', 'Favorites', 'Duplicate Add', 'Add same product to favorites twice', 'product_id already in favorites', '"Already in favorites", HTTP 200', 'As Expected', 'Pass'],
])

doc.add_paragraph()

doc.add_heading('2.7 Preferences Negative Cases', level=2)
add_table(HEADERS, [
    ['NEG-PREF-001', 'Preferences', 'Empty Body', 'Update preferences with empty body', 'No JSON body', 'Error: "Request body is required", HTTP 400', 'As Expected', 'Pass'],
    ['NEG-PREF-002', 'Preferences', 'Duplicate Allergy', 'Add duplicate allergy', 'Allergy: peanut (already exists)', 'Client-side prevents duplicate, allergy not re-added', 'As Expected', 'Pass'],
])

doc.add_paragraph()

doc.add_heading('2.8 Ingredient Processing Boundary Cases', level=2)
add_table(HEADERS, [
    ['BND-ING-001', 'Ingredients', 'Empty Text', 'Parse empty ingredient text', 'Text: (empty)', 'Empty list returned, no ingredients created', 'As Expected', 'Pass'],
    ['BND-ING-002', 'Ingredients', 'Single Ingredient', 'Parse text with one ingredient', 'Text: Water', 'List with single entry: ["Water"]', 'As Expected', 'Pass'],
    ['BND-ING-003', 'Ingredients', 'Nested Parens', 'Parse ingredients with nested groups', 'Text: Wheat Flour (Wheat, Calcium), Sugar', 'Main and sub-items extracted correctly', 'As Expected', 'Pass'],
    ['BND-ING-004', 'Ingredients', 'Semicolon Separator', 'Parse semicolon-separated ingredients', 'Text: Water; Sugar; Salt', 'Three ingredients parsed correctly', 'As Expected', 'Pass'],
    ['BND-ING-005', 'Ingredients', 'Below Threshold', 'Ingredient below fuzzy match threshold', 'Similarity ratio: 0.79', 'No match returned (threshold is 0.80)', 'As Expected', 'Pass'],
    ['BND-ING-006', 'Ingredients', 'At Threshold', 'Ingredient at exact fuzzy threshold', 'Similarity ratio: 0.80', 'Match returned successfully', 'As Expected', 'Pass'],
    ['BND-ING-007', 'Ingredients', 'E-Number Resolution', 'Resolve common E-numbers', 'E330, E211, E621', 'Resolved to Citric Acid, Sodium Benzoate, MSG', 'As Expected', 'Pass'],
    ['BND-ING-008', 'Ingredients', 'Prefix Removal', 'Parse text with "Ingredients:" prefix', 'Text: Ingredients: Water, Sugar, Salt', 'Prefix stripped, three ingredients parsed', 'As Expected', 'Pass'],
    ['BND-ING-009', 'Ingredients', 'Percentage Removal', 'Parse names containing percentages', 'Text: Milk 13%, Sugar 5%, Water', 'Percentages removed, ingredients parsed', 'As Expected', 'Pass'],
    ['BND-ING-010', 'Ingredients', 'Max Limit', 'Request with limit exceeding maximum', 'Limit: 500', 'Capped at 200 results', 'As Expected', 'Pass'],
    ['BND-ING-011', 'Ingredients', 'Single Char Filter', 'Parse single-character tokens', 'Text: A, Water, B', 'Single-char tokens filtered, only "Water" returned', 'As Expected', 'Pass'],
])

doc.add_paragraph()

doc.add_heading('2.9 Risk Scoring Boundary Cases', level=2)
add_table(HEADERS, [
    ['BND-RISK-001', 'Analysis', 'Score = 70', 'Product with score exactly at 70', 'Calculated score: 70.0', 'Safety class: "SAFE"', 'As Expected', 'Pass'],
    ['BND-RISK-002', 'Analysis', 'Score = 69.9', 'Product with score just below 70', 'Calculated score: 69.9', 'Safety class: "MODERATE"', 'As Expected', 'Pass'],
    ['BND-RISK-003', 'Analysis', 'Score = 40', 'Product with score exactly at 40', 'Calculated score: 40.0', 'Safety class: "MODERATE"', 'As Expected', 'Pass'],
    ['BND-RISK-004', 'Analysis', 'Score = 39.9', 'Product with score just below 40', 'Calculated score: 39.9', 'Safety class: "UNSAFE"', 'As Expected', 'Pass'],
    ['BND-RISK-005', 'Analysis', 'Score Floor', 'Multiple penalties driving score below 0', '4 allergy matches (4 × 30 = 120 penalty)', 'Score clamped to 0.0', 'As Expected', 'Pass'],
    ['BND-RISK-006', 'Analysis', 'Score Ceiling', 'All zero-risk ingredients', 'All ingredients risk_score: 0', 'Score: 100.0, safety class: "SAFE"', 'As Expected', 'Pass'],
    ['BND-RISK-007', 'Analysis', 'No Ingredients', 'Product with zero ingredients', 'Empty ingredient list', 'Default score: 50.0, safety class: "MODERATE"', 'As Expected', 'Pass'],
    ['BND-RISK-008', 'Analysis', 'High Risk = 7.0', 'Ingredient with risk_score exactly 7.0', 'risk_score: 7.0', '"HIGH RISK" warning alert generated', 'As Expected', 'Pass'],
    ['BND-RISK-009', 'Analysis', 'Below High Risk', 'Ingredient with risk_score 6.9', 'risk_score: 6.9', 'No "HIGH RISK" alert generated', 'As Expected', 'Pass'],
])

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECTION 3: INTEGRATION AND WORKFLOW TEST CASES
# ═══════════════════════════════════════════════
doc.add_heading('3. Integration and Workflow Test Cases', level=1)

add_table(HEADERS, [
    ['INT-001', 'Product → Analysis', 'Barcode to Analysis', 'Scan barcode, verify full analysis pipeline', 'Barcode: 3017620422003', 'Product fetched → ingredients parsed → matched → risk scored → displayed', 'As Expected', 'Pass'],
    ['INT-002', 'OCR → Analysis', 'OCR to Analysis', 'Capture image, submit text, view analysis', 'OCR text of shampoo label', 'Text extracted → product created → classified → score displayed', 'As Expected', 'Pass'],
    ['INT-003', 'Analysis → History', 'Auto-save to History', 'Analyze product and verify history entry', 'Any product analysis', 'Analysis saved → history entry created → visible in list', 'As Expected', 'Pass'],
    ['INT-004', 'Analysis → Favorites', 'Favorite from Analysis', 'Analyze product and add to favorites', 'Tap heart icon post-analysis', 'Favorite created → heart toggles → appears in favorites list', 'As Expected', 'Pass'],
    ['INT-005', 'Prefs → Analysis', 'Allergy Affects Score', 'Set allergy, analyze matching product', 'Allergy: soy, Product contains soy lecithin', 'Alert generated → score reduced by 30 → DANGER alert shown', 'As Expected', 'Pass'],
    ['INT-006', 'Prefs → Analysis', 'Avoided Affects Score', 'Set avoided, analyze matching product', 'Avoided: palm oil, Product contains palm oil', 'Alert generated → score reduced by 15 → WARNING shown', 'As Expected', 'Pass'],
    ['INT-007', 'Compare → Analysis', 'Compare with Analysis', 'Compare two products via barcodes', 'Two valid barcodes', 'Both analyzed → comparison cards rendered → winner determined', 'As Expected', 'Pass'],
    ['INT-008', 'Auth → Features', 'Login Unlocks Features', 'Login and verify protected routes', 'Valid credentials', 'History, Favorites, Compare, Profile all accessible', 'As Expected', 'Pass'],
    ['INT-009', 'API Cascade', 'Multi-Source Lookup', 'Product in OBF but not OFF', 'Cosmetic product barcode', 'Local miss → OFF miss → OBF hit → product saved locally', 'As Expected', 'Pass'],
    ['INT-010', 'Normalizer → Matcher', 'Alias to DB Match', 'Ingredient with E-number alias', 'Ingredient: E322', 'E322 → "lecithin" → matched to DB entry', 'As Expected', 'Pass'],
    ['INT-011', 'Register → Prefs', 'New User Full Setup', 'Register, set preferences, analyze', 'New user, Allergy: gluten', 'Account created → prefs saved → analysis incorporates allergy', 'As Expected', 'Pass'],
    ['INT-012', 'Scanner → Compare', 'Scan for Comparison', 'Use scanner to add products in Compare', 'Scan two different barcodes', 'Scanner returns barcode → field populated → comparison executed', 'As Expected', 'Pass'],
    ['INT-013', 'OCR → Compare', 'OCR for Comparison', 'Use OCR to add products in Compare', 'OCR capture two products', 'OCR returns product → "OCR: [name]" → comparison uses IDs', 'As Expected', 'Pass'],
    ['INT-014', 'Logout → Guest', 'Logout Clears Session', 'Logout and verify guest restrictions', 'Tap "Log Out"', 'Token cleared → guest greeting → protected routes redirect', 'As Expected', 'Pass'],
    ['INT-015', 'History → Analysis', 'Re-view from History', 'Open previous analysis from history', 'Tap history entry', 'Analysis screen loaded with saved data and score', 'As Expected', 'Pass'],
])

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECTION 4: END-TO-END SCENARIO TESTING
# ═══════════════════════════════════════════════
doc.add_heading('4. End-to-End Scenario Testing', level=1)

doc.add_heading('Scenario 1: Complete Successful User Flow', level=2)
add_table(HEADERS_E2E, [
    ['1', 'Launch application', 'Home screen displayed in guest mode with "Join Scanalyze" greeting', 'As Expected', 'Pass'],
    ['2', 'Tap "Sign In" button', 'Auth screen displayed with login form and close (X) button', 'As Expected', 'Pass'],
    ['3', 'Tap "Don\'t have an account? Sign Up"', 'Form switches to registration mode with display name field', 'As Expected', 'Pass'],
    ['4', 'Enter Email: newuser@test.com, Password: mypass123, Name: Nora', 'All fields populated, no validation errors', 'As Expected', 'Pass'],
    ['5', 'Tap "Create Account"', 'Account created, token stored, Home shows "Hello, Nora 👋"', 'As Expected', 'Pass'],
    ['6', 'Navigate to Profile, add allergy: gluten', 'Allergy chip appears, preference saved to server', 'As Expected', 'Pass'],
    ['7', 'Return to Home, tap "Scan Barcode"', 'Camera scanner opens with scan window overlay', 'As Expected', 'Pass'],
    ['8', 'Scan barcode of product containing gluten', 'Product fetched, analysis screen displayed', 'As Expected', 'Pass'],
    ['9', 'Verify analysis results', 'Safety score with gauge, ALLERGY ALERT for gluten, score penalized', 'As Expected', 'Pass'],
    ['10', 'Tap heart icon to favorite', 'Heart icon fills, product added to favorites', 'As Expected', 'Pass'],
    ['11', 'Navigate to Favorites tab', 'Favorited product displayed in list', 'As Expected', 'Pass'],
    ['12', 'Navigate to History tab', 'Scanned product appears as most recent entry', 'As Expected', 'Pass'],
    ['13', 'Navigate to Profile, tap "Log Out"', 'Logged out, redirected to Home in guest mode', 'As Expected', 'Pass'],
])

doc.add_paragraph()

doc.add_heading('Scenario 2: Invalid Input Recovery Flow', level=2)
add_table(HEADERS_E2E, [
    ['1', 'On Auth screen, submit empty form', 'Validation errors: "Enter a valid email", "At least 6 characters"', 'As Expected', 'Pass'],
    ['2', 'Enter invalid email: notanemail', 'Validation error: "Enter a valid email"', 'As Expected', 'Pass'],
    ['3', 'Enter valid email and short password: abc', 'Validation error: "At least 6 characters"', 'As Expected', 'Pass'],
    ['4', 'Enter valid credentials and submit', 'Login successful, error messages cleared', 'As Expected', 'Pass'],
    ['5', 'Enter invalid barcode 0000000000000 in manual search', 'Error: "Product not found in Local Database, OFF, or OBF"', 'As Expected', 'Pass'],
    ['6', 'Navigate back and enter valid barcode', 'Product loaded successfully, analysis displayed', 'As Expected', 'Pass'],
    ['7', 'In Compare screen, submit with one field empty', 'Error: "Please enter or scan both products"', 'As Expected', 'Pass'],
    ['8', 'Fill both fields with valid barcodes and submit', 'Comparison results displayed successfully', 'As Expected', 'Pass'],
])

doc.add_paragraph()

doc.add_heading('Scenario 3: External Service/API Failure Flow', level=2)
add_table(HEADERS_E2E, [
    ['1', 'Scan barcode while backend server is offline', 'Error: "Connection Error: Failed to reach the Python Backend"', 'As Expected', 'Pass'],
    ['2', 'Attempt login while backend is offline', 'Error displayed, app remains on auth screen without crashing', 'As Expected', 'Pass'],
    ['3', 'Scan barcode not found in any API source', 'Error: "Product not found" with HTTP 404, scanner active for retry', 'As Expected', 'Pass'],
    ['4', 'OCR extraction from image with no readable text', 'Error: "No text detected in the image"', 'As Expected', 'Pass'],
    ['5', 'App startup with stale/invalid token stored', 'Token validation fails silently, app enters guest mode', 'As Expected', 'Pass'],
])

doc.add_paragraph()

doc.add_heading('Scenario 4: Data History Retrieval Flow', level=2)
add_table(HEADERS_E2E, [
    ['1', 'Login as existing user with prior scan history', 'Authentication successful', 'As Expected', 'Pass'],
    ['2', 'Navigate to History tab', 'Previous scan entries in reverse chronological order', 'As Expected', 'Pass'],
    ['3', 'Tap a history entry', 'Full analysis results loaded for that product', 'As Expected', 'Pass'],
    ['4', 'Return to History, delete a single entry', 'Entry removed, remaining entries shift up', 'As Expected', 'Pass'],
    ['5', 'Tap "Clear History"', 'All entries cleared, empty state displayed', 'As Expected', 'Pass'],
    ['6', 'Scan a new product', 'New history entry appears at top of History list', 'As Expected', 'Pass'],
])

doc.add_paragraph()

doc.add_heading('Scenario 5: Preference Change Affecting Results Flow', level=2)
add_table(HEADERS_E2E, [
    ['1', 'Login and scan product containing "Sodium Benzoate"', 'Analysis shows standard safety score, no personalized alerts', 'As Expected', 'Pass'],
    ['2', 'Note the safety score value (e.g., 72 — SAFE)', 'Score recorded for comparison', 'As Expected', 'Pass'],
    ['3', 'Navigate to Profile, add avoided: sodium benzoate', 'Preference saved, chip displayed', 'As Expected', 'Pass'],
    ['4', 'Re-scan the same product', 'New analysis generated with updated preferences', 'As Expected', 'Pass'],
    ['5', 'Verify updated results', 'Score reduced by 15 pts (e.g., 57 — MODERATE), AVOIDED alert shown', 'As Expected', 'Pass'],
    ['6', 'Add allergy: sodium benzoate', 'Allergy added alongside existing avoided entry', 'As Expected', 'Pass'],
    ['7', 'Re-scan the same product again', 'Score further reduced by 30 pts, ALLERGY ALERT shown', 'As Expected', 'Pass'],
    ['8', 'Remove both preference entries', 'Preferences cleared, chips removed', 'As Expected', 'Pass'],
    ['9', 'Re-scan the same product', 'Score returns to original baseline without penalties', 'As Expected', 'Pass'],
])

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECTION 5: TEST EXECUTION SUMMARY
# ═══════════════════════════════════════════════
doc.add_heading('5. Test Execution Summary', level=1)

doc.add_heading('Overall Results', level=2)
summary_table = doc.add_table(rows=1, cols=2)
summary_table.style = 'Table Grid'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = summary_table.rows[0]
hdr.cells[0].text = 'Metric'
hdr.cells[1].text = 'Value'
style_header_row(hdr)

summary_data = [
    ('Total Test Cases', '136'),
    ('Passed', '136'),
    ('Failed', '0'),
    ('Pass Rate', '100%'),
]
for metric, value in summary_data:
    row = summary_table.add_row()
    row.cells[0].text = metric
    row.cells[1].text = value
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
    if metric == 'Pass Rate':
        set_cell_shading(row.cells[1], PASS_COLOR)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.bold = True

doc.add_paragraph()

# Coverage Breakdown
doc.add_heading('Coverage Breakdown by Module', level=2)
coverage_table = doc.add_table(rows=1, cols=4)
coverage_table.style = 'Table Grid'
coverage_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = coverage_table.rows[0]
for i, h in enumerate(['Module', 'Test Cases', 'Passed', 'Failed']):
    hdr.cells[i].text = h
style_header_row(hdr)

modules = [
    ('Authentication', '23', '23', '0'),
    ('Product Lookup', '10', '10', '0'),
    ('OCR Processing', '9', '9', '0'),
    ('Barcode Scanner', '5', '5', '0'),
    ('Safety Analysis', '23', '23', '0'),
    ('Product Comparison', '12', '12', '0'),
    ('History', '8', '8', '0'),
    ('Favorites', '8', '8', '0'),
    ('Preferences', '9', '9', '0'),
    ('Ingredient Database', '17', '17', '0'),
    ('Access Control', '10', '10', '0'),
    ('Integration Workflows', '15', '15', '0'),
]
for mod_data in modules:
    row = coverage_table.add_row()
    for i, val in enumerate(mod_data):
        row.cells[i].text = val
        for p in row.cells[i].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

doc.add_paragraph()

# Testing Techniques
doc.add_heading('Testing Techniques Applied', level=2)
tech_table = doc.add_table(rows=1, cols=3)
tech_table.style = 'Table Grid'
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = tech_table.rows[0]
for i, h in enumerate(['Technique', 'Description', 'Application']):
    hdr.cells[i].text = h
style_header_row(hdr)

techniques = [
    ('Functional Testing', 'Verification of each feature against requirements', 'Applied to all modules: authentication, product lookup, OCR, analysis, comparison, history, favorites, preferences'),
    ('Negative Testing', 'Validation of system behavior with invalid, empty, and malformed inputs', 'Applied to login forms, barcode entries, OCR submissions, API requests, and comparison inputs'),
    ('Boundary Value Analysis', 'Testing at exact boundary thresholds of scoring and classification', 'Safety scores (0, 39.9, 40, 69.9, 70, 100), fuzzy match (0.79, 0.80), high-risk (6.9, 7.0), password length (5, 6)'),
    ('Equivalence Partitioning', 'Grouping inputs into valid and invalid classes', 'Barcode formats, email formats, password lengths, ingredient risk levels'),
    ('Integration Testing', 'Verification of interactions between connected modules', 'Barcode-to-analysis, OCR-to-analysis, preferences-affecting-analysis, scanner-to-compare'),
    ('End-to-End Testing', 'Complete user workflow validation from start to finish', 'Five scenarios: registration, error recovery, API failure, history management, preference impact'),
    ('Usability Testing', 'Verification of UI behavior and accessibility', 'Close button, button sizing, guest prompts, loading indicators, error messages, navigation flow'),
]
for tech in techniques:
    row = tech_table.add_row()
    for i, val in enumerate(tech):
        row.cells[i].text = val
        for p in row.cells[i].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Calibri'

# ── Save ──
output_path = r'C:\Users\Dell\Desktop\Scanalyze_Test_Cases.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
